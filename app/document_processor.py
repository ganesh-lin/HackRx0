import PyPDF2
import requests
import os
import logging
import re
import tempfile
import time
import uuid
from urllib.parse import urlparse
from typing import List, Dict, Any
from docx import Document
import email
from email.mime.text import MIMEText
from bs4 import BeautifulSoup
try:
    import tiktoken
    TIKTOKEN_AVAILABLE = True
except ImportError:
    TIKTOKEN_AVAILABLE = False
from dotenv import load_dotenv

load_dotenv()

logging.basicConfig(filename="app.log", level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

class DocumentProcessor:
    def __init__(self):
        self.max_chunk_size = int(os.getenv("MAX_CHUNK_SIZE", 512))
        self.chunk_overlap = int(os.getenv("CHUNK_OVERLAP", 50))
        if TIKTOKEN_AVAILABLE:
            self.encoding = tiktoken.get_encoding("cl100k_base")
        else:
            self.encoding = None
    
    def download_document(self, url: str) -> str:
        """Download document from URL and save locally."""
        try:
            response = requests.get(url, timeout=30, stream=True)
            response.raise_for_status()
            
            # Determine file extension from URL or content type
            parsed_url = urlparse(url)
            file_extension = os.path.splitext(parsed_url.path)[1].lower()
            
            if not file_extension:
                content_type = response.headers.get('content-type', '').lower()
                if 'pdf' in content_type:
                    file_extension = '.pdf'
                elif 'docx' in content_type or 'document' in content_type:
                    file_extension = '.docx'
                else:
                    file_extension = '.txt'
            
            # Create temporary file with unique name
            unique_id = str(uuid.uuid4())[:8]
            timestamp = str(int(time.time()))
            temp_filename = f"hackrx_doc_{timestamp}_{unique_id}{file_extension}"
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension, prefix=temp_filename) as temp_file:
                for chunk in response.iter_content(chunk_size=8192):
                    temp_file.write(chunk)
                temp_path = temp_file.name
            
            logging.info(f"Downloaded document from {url} to {temp_path}")
            return temp_path
            
        except Exception as e:
            logging.error(f"Failed to download document: {str(e)}")
            raise

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF with improved error handling."""
        text = ""
        try:
            with open(pdf_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num + 1} ---\n"
                            text += page_text
                    except Exception as e:
                        logging.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
                
                logging.info(f"Extracted text from PDF: {len(text)} characters")
                
        except Exception as e:
            logging.error(f"Failed to extract text from PDF: {str(e)}")
            raise
        finally:
            # Ensure file is closed and removed
            try:
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
            except Exception as e:
                logging.warning(f"Failed to remove temporary file {pdf_path}: {e}")
        
        return self.clean_text(text)

    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file."""
        text = ""
        try:
            doc = Document(docx_path)
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
                
            logging.info(f"Extracted text from DOCX: {len(text)} characters")
                
        except Exception as e:
            logging.error(f"Failed to extract text from DOCX: {str(e)}")
            raise
        finally:
            # Ensure file is closed and removed
            try:
                if os.path.exists(docx_path):
                    os.remove(docx_path)
            except Exception as e:
                logging.warning(f"Failed to remove temporary file {docx_path}: {e}")
        
        return self.clean_text(text)

    def extract_text_from_email(self, email_path: str) -> str:
        """Extract text from email file."""
        try:
            with open(email_path, 'r', encoding='utf-8') as file:
                msg = email.message_from_file(file)
            
            text = ""
            
            # Extract headers
            text += f"Subject: {msg.get('Subject', 'N/A')}\n"
            text += f"From: {msg.get('From', 'N/A')}\n"
            text += f"To: {msg.get('To', 'N/A')}\n"
            text += f"Date: {msg.get('Date', 'N/A')}\n\n"
            
            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        text += part.get_payload(decode=True).decode('utf-8', errors='ignore')
                    elif part.get_content_type() == "text/html":
                        html_content = part.get_payload(decode=True).decode('utf-8', errors='ignore')
                        soup = BeautifulSoup(html_content, 'html.parser')
                        text += soup.get_text()
            else:
                if msg.get_content_type() == "text/plain":
                    text += msg.get_payload(decode=True).decode('utf-8', errors='ignore')
                elif msg.get_content_type() == "text/html":
                    html_content = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
                    soup = BeautifulSoup(html_content, 'html.parser')
                    text += soup.get_text()
            
            # Clean up temporary file
            if os.path.exists(email_path):
                os.remove(email_path)
                
            logging.info(f"Extracted text from email: {len(text)} characters")
            return self.clean_text(text)
            
        except Exception as e:
            logging.error(f"Failed to extract text from email: {str(e)}")
            raise

    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}\"\'\/\\\&\%\$\#\@]', ' ', text)
        
        # Remove excessive spaces
        text = re.sub(r' +', ' ', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text

    def process_document(self, url: str) -> str:
        """Main method to process any document type."""
        try:
            file_path = self.download_document(url)
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                return self.extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                return self.extract_text_from_docx(file_path)
            elif file_extension in ['.eml', '.msg']:
                return self.extract_text_from_email(file_path)
            else:
                # Try to read as plain text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                os.remove(file_path)
                return self.clean_text(text)
                
        except Exception as e:
            logging.error(f"Failed to process document: {str(e)}")
            raise

    def chunk_text(self, text: str) -> List[Dict[str, Any]]:
        """Split text into meaningful chunks with enhanced overlap for insurance documents."""
        try:
            # Clean text and normalize whitespace
            text = re.sub(r'\s+', ' ', text)
            
            # First try to identify common insurance policy sections
            section_patterns = [
                r'(?i)(Definitions|Defined\s+Terms)\s*[:\n]',
                r'(?i)(Waiting\s+Period|Waiting\s+Periods)\s*[:\n]',
                r'(?i)(Benefits|Coverage|Coverages)\s*[:\n]',
                r'(?i)(Exclusions|Limitations)\s*[:\n]',
                r'(?i)(Terms\s+and\s+Conditions)\s*[:\n]',
                r'(?i)(Premium\s+Payment|Grace\s+Period)\s*[:\n]',
                r'(?i)(Renewal\s+Conditions|Portability)\s*[:\n]',
                r'(?i)(Claim\s+Process|Claim\s+Procedure)\s*[:\n]',
                r'(?i)(Cancellation|Termination)\s*[:\n]',
                r'(?i)(Maternity\s+Benefits?)\s*[:\n]',
                r'(?i)(Organ\s+Donor\s+Expenses?)\s*[:\n]',
                r'(?i)(Pre-existing\s+Disease|PED)\s*[:\n]',
                r'(?i)(Sub-limits|Room\s+Rent\s+Capping)\s*[:\n]',
                r'(?i)(AYUSH\s+Coverage|AYUSH\s+Treatment)\s*[:\n]',
                r'(?i)(No\s+Claim\s+Bonus|NCB)\s*[:\n]',
                r'(?i)(Preventive\s+Health\s+Check[-\s]?up)\s*[:\n]',
                r'(?i)(Physiotherapy|Domiciliary\s+Treatment)\s*[:\n]',
                r'(?i)(Specific\s+Disease\s+Waiting\s+Period)\s*[:\n]',
                r'(?i)(Initial\s+Waiting\s+Period|Cooling\s+Off\s+Period)\s*[:\n]',
                r'(?i)(Hospital\s+Cash|Daily\s+Cash\s+Allowance)\s*[:\n]'
            ]
            
            # Find all section matches
            section_matches = []
            for pattern in section_patterns:
                for match in re.finditer(pattern, text):
                    section_matches.append((match.start(), match.group()))
            
            # Sort by position in text
            section_matches.sort(key=lambda x: x[0])
            
            # If we found meaningful sections, use them for chunking
            chunks = []
            chunk_id = 0
            
            if len(section_matches) >= 3:  # Only use section-based chunking if we found enough sections
                for i in range(len(section_matches)):
                    start_pos = section_matches[i][0]
                    
                    # Determine end position (either next section or end of text)
                    if i < len(section_matches) - 1:
                        end_pos = section_matches[i+1][0]
                        # Include some overlap (200 chars) with next section
                        end_pos = min(end_pos + 200, len(text))
                    else:
                        end_pos = len(text)
                    
                    section_text = text[start_pos:end_pos].strip()
                    
                    # Further chunk large sections
                    if len(section_text) > self.max_chunk_size * 2:  # Very large section
                        sub_chunks = self._chunk_large_text(section_text, chunk_id)
                        chunks.extend(sub_chunks)
                        chunk_id += len(sub_chunks)
                    else:
                        # Calculate tokens
                        if self.encoding:
                            tokens = len(self.encoding.encode(section_text))
                        else:
                            tokens = len(section_text.split()) * 1.3  # Rough approximation
                        
                        chunks.append({
                            "id": chunk_id,
                            "text": section_text,
                            "tokens": tokens,
                            "start_sentence": chunk_id * 10,  # Approximate
                            "end_sentence": chunk_id * 10 + len(section_text.split('.'))
                        })
                        chunk_id += 1
            else:
                # Fall back to improved sentence-based chunking
                # Split text into sentences with better handling
                sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)
                
                current_chunk = ""
                current_tokens = 0
                
                for sentence in sentences:
                    if not sentence.strip():
                        continue
                        
                    sentence = sentence.strip()
                    
                    # Calculate tokens - fallback to word count if tiktoken not available
                    if self.encoding:
                        sentence_tokens = len(self.encoding.encode(sentence))
                    else:
                        sentence_tokens = len(sentence.split()) * 1.3  # Rough approximation
                    
                    # If adding this sentence would exceed max chunk size
                    if current_tokens + sentence_tokens > self.max_chunk_size and current_chunk:
                        # Save current chunk
                        chunks.append({
                            "id": chunk_id,
                            "text": current_chunk.strip(),
                            "tokens": current_tokens,
                            "start_sentence": chunk_id * 10,  # Approximate
                            "end_sentence": chunk_id * 10 + len(current_chunk.split('.'))
                        })
                        
                        # Start new chunk with more overlap for insurance documents
                        overlap_text = ""
                        overlap_tokens = 0
                        
                        # Add overlap from end of previous chunk - increased overlap for insurance
                        current_sentences = current_chunk.split('.')
                        for i in range(min(5, len(current_sentences))):  # Take last 5 sentences for overlap
                            overlap_sentence = current_sentences[-(i+1)].strip()
                            if overlap_sentence:
                                if self.encoding:
                                    overlap_sentence_tokens = len(self.encoding.encode(overlap_sentence))
                                else:
                                    overlap_sentence_tokens = len(overlap_sentence.split()) * 1.3
                                
                                if overlap_tokens + overlap_sentence_tokens <= self.chunk_overlap * 1.5:  # Increased overlap
                                    overlap_text = overlap_sentence + ". " + overlap_text
                                    overlap_tokens += overlap_sentence_tokens
                                else:
                                    break
                        
                        current_chunk = overlap_text + sentence + ". "
                        current_tokens = overlap_tokens + sentence_tokens
                        chunk_id += 1
                    else:
                        current_chunk += sentence + ". "
                        current_tokens += sentence_tokens
                
                # Add the last chunk if it has content
                if current_chunk.strip():
                    chunks.append({
                        "id": chunk_id,
                        "text": current_chunk.strip(),
                        "tokens": current_tokens,
                        "start_sentence": chunk_id * 10,
                        "end_sentence": chunk_id * 10 + len(current_chunk.split('.'))
                    })
            
            # Post-process: ensure critical insurance terms are fully captured
            # These sections should never be split across chunks if possible
            final_chunks = self._optimize_insurance_chunks(chunks)
            
            logging.info(f"Created {len(final_chunks)} chunks from text")
            return final_chunks
            
        except Exception as e:
            logging.error(f"Failed to chunk text: {str(e)}")
            # Fallback: simple splitting
            chunk_size = 1000
            chunks = []
            for i in range(0, len(text), chunk_size):
                chunk_text = text[i:i + chunk_size]
                tokens = len(chunk_text.split()) if not self.encoding else len(self.encoding.encode(chunk_text))
                chunks.append({
                    "id": i // chunk_size,
                    "text": chunk_text,
                    "tokens": tokens,
                    "start_sentence": i,
                    "end_sentence": i + chunk_size
                })
            return chunks
        except Exception as e:
            logging.error(f"Failed to chunk text: {str(e)}")
            # Fallback: simple splitting
            chunk_size = 1000
            chunks = []
            for i in range(0, len(text), chunk_size):
                chunk_text = text[i:i + chunk_size]
                tokens = len(chunk_text.split()) if not self.encoding else len(self.encoding.encode(chunk_text))
                chunks.append({
                    "id": i // chunk_size,
                    "text": chunk_text,
                    "tokens": tokens,
                    "start_sentence": i,
                    "end_sentence": i + chunk_size
                })
            return chunks
            
    def _chunk_large_text(self, text: str, start_id: int) -> List[Dict[str, Any]]:
        """Further chunk large sections of text."""
        # Split into paragraphs
        paragraphs = text.split('\n\n')
        if len(paragraphs) == 1:  # If no clear paragraphs, split by period
            paragraphs = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)
        
        chunks = []
        current_chunk = ""
        current_tokens = 0
        chunk_id = start_id
        
        for para in paragraphs:
            if not para.strip():
                continue
                
            para = para.strip()
            
            # Calculate tokens
            if self.encoding:
                para_tokens = len(self.encoding.encode(para))
            else:
                para_tokens = len(para.split()) * 1.3
            
            # If adding this paragraph would exceed max chunk size
            if current_tokens + para_tokens > self.max_chunk_size and current_chunk:
                # Save current chunk
                chunks.append({
                    "id": chunk_id,
                    "text": current_chunk.strip(),
                    "tokens": current_tokens,
                    "start_sentence": chunk_id * 10,
                    "end_sentence": chunk_id * 10 + len(current_chunk.split('.'))
                })
                
                current_chunk = para
                current_tokens = para_tokens
                chunk_id += 1
            else:
                if current_chunk:
                    current_chunk += " " + para
                else:
                    current_chunk = para
                current_tokens += para_tokens
        
        # Add the last chunk if it has content
        if current_chunk.strip():
            chunks.append({
                "id": chunk_id,
                "text": current_chunk.strip(),
                "tokens": current_tokens,
                "start_sentence": chunk_id * 10,
                "end_sentence": chunk_id * 10 + len(current_chunk.split('.'))
            })
        
        return chunks
    
    def _optimize_insurance_chunks(self, chunks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Optimize chunks for insurance document retrieval."""
        # Ensure critical insurance terms are fully captured in chunks
        critical_terms = [
            (r'(?i)grace\s+period', 'grace period'),
            (r'(?i)waiting\s+period\s+for\s+pre-existing', 'pre-existing disease waiting period'),
            (r'(?i)maternity\s+expenses?', 'maternity coverage'),
            (r'(?i)cataract\s+surgery', 'cataract surgery waiting period'),
            (r'(?i)organ\s+donor', 'organ donor coverage'),
            (r'(?i)no\s+claim\s+discount', 'no claim discount'),
            (r'(?i)no\s+claim\s+bonus', 'no claim bonus'),
            (r'(?i)health\s+check-up', 'health check-up benefit'),
            (r'(?i)definition\s+of\s+hospital', 'hospital definition'),
            (r'(?i)AYUSH\s+treatment', 'AYUSH coverage'),
            (r'(?i)room\s+rent\s+capping', 'room rent sublimits'),
            (r'(?i)sub-limits', 'sublimits'),
            (r'(?i)physiotherapy', 'physiotherapy coverage')
        ]
        
        # Create mapping of chunks containing critical terms
        term_to_chunks = {}
        for i, chunk in enumerate(chunks):
            for pattern, term_name in critical_terms:
                if re.search(pattern, chunk["text"], re.IGNORECASE):
                    if term_name not in term_to_chunks:
                        term_to_chunks[term_name] = []
                    term_to_chunks[term_name].append(i)
        
        # Merge small chunks to ensure they meet the minimum size and contain full contexts
        final_chunks = []
        i = 0
        while i < len(chunks):
            current = chunks[i]
            
            # Find if current chunk contains start of a critical term but not its entirety
            for term_name, chunk_indices in term_to_chunks.items():
                if i in chunk_indices and i+1 < len(chunks) and i+1 in chunk_indices:
                    # Merge current with next to keep term together
                    next_chunk = chunks[i+1]
                    merged_text = current["text"] + " " + next_chunk["text"]
                    
                    # Calculate tokens
                    if self.encoding:
                        merged_tokens = len(self.encoding.encode(merged_text))
                    else:
                        merged_tokens = len(merged_text.split()) * 1.3
                    
                    if merged_tokens <= self.max_chunk_size * 1.5:  # Allow slightly larger chunks for critical terms
                        merged_chunk = {
                            "id": current["id"],
                            "text": merged_text,
                            "tokens": merged_tokens,
                            "start_sentence": current["start_sentence"],
                            "end_sentence": next_chunk["end_sentence"],
                            "critical_term": term_name  # Mark that this contains a critical term
                        }
                        final_chunks.append(merged_chunk)
                        i += 2  # Skip next chunk as it's merged
                        break
            else:
                # No critical term spanning across chunks, process normally
                if i < len(chunks) - 1 and current["tokens"] < 100:
                    # Merge very small chunks
                    next_chunk = chunks[i+1]
                    if current["tokens"] + next_chunk["tokens"] <= self.max_chunk_size:
                        merged_chunk = {
                            "id": current["id"],
                            "text": current["text"] + " " + next_chunk["text"],
                            "tokens": current["tokens"] + next_chunk["tokens"],
                            "start_sentence": current["start_sentence"],
                            "end_sentence": next_chunk["end_sentence"]
                        }
                        final_chunks.append(merged_chunk)
                        i += 2
                        continue
                
                final_chunks.append(current)
                i += 1
        
        return final_chunks

# Backward compatibility functions
def download_pdf(url: str) -> str:
    """Backward compatibility function."""
    processor = DocumentProcessor()
    return processor.download_document(url)

def extract_text_from_pdf(pdf_path: str) -> str:
    """Backward compatibility function."""
    processor = DocumentProcessor()
    return processor.extract_text_from_pdf(pdf_path)